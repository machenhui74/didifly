import os
# 延迟导入其他模块
# import pandas as pd
# import random
# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
# import docx
from tkinter import messagebox

class TrainingPlanGenerator:
    def __init__(self):
        pass
    
    def get_age_range(self, age):
        """根据年龄获取对应的年龄范围分类"""
        try:
            age = int(age)
            if 4 <= age <= 6:
                return '4-6'
            elif 7 <= age <= 9:
                return '7-9'
            elif 10 <= age <= 12:
                return '10-12'
            else:
                return None
        except ValueError:
            return None
            
    def generate_training_plan(self, name, age, selected_tags, action_db_path, output_dir):
        """生成训练方案，并保存为Word文档"""
        try:
            # 按需导入模块
            import pandas as pd
            import random
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import docx
            
            # 确定年龄范围
            age_range = self.get_age_range(age)
            if not age_range:
                messagebox.showwarning("警告", "无法确定年龄范围，请确保年龄在4-12岁之间")
                return
            
            # 读取动作库
            df = pd.read_excel(action_db_path)
            
            # 过滤掉动作名称为空的行
            df = df.dropna(subset=['动作名称']).reset_index(drop=True)
            
            # 处理标签匹配逻辑 - 支持#分隔的多标签
            # 筛选年龄范围符合的动作
            age_filtered_df = df[df['年龄范围'] == age_range]
            
            if age_filtered_df.empty:
                messagebox.showwarning("警告", f"在动作库中未找到适合{age}岁的训练动作")
                return
            
            # 计算标签匹配分数，优先选择标签匹配度高的动作
            def calculate_match_score(tag_str):
                if not isinstance(tag_str, str):
                    return 0
                
                action_tags = [tag.strip() for tag in tag_str.split('#') if tag.strip()]
                if not action_tags:
                    return 0
                
                # 计算用户选择标签与动作标签的交集
                matching_tags = set(action_tags).intersection(set(selected_tags))
                
                # 计算匹配分数
                # 1. 完全匹配情况：用户选择标签与动作标签完全相同
                if set(action_tags) == set(selected_tags):
                    return 100
                
                # 2. 用户只选了一个标签，动作也只有这一个标签
                if len(selected_tags) == 1 and len(action_tags) == 1 and action_tags[0] in selected_tags:
                    return 90
                
                # 3. 匹配标签占动作标签的比例
                action_match_ratio = len(matching_tags) / len(action_tags)
                
                # 4. 匹配标签占用户选择标签的比例
                user_match_ratio = len(matching_tags) / len(selected_tags)
                
                # 综合评分 (加权平均)
                return (action_match_ratio * 30) + (user_match_ratio * 30) + (len(matching_tags) * 10)
            
            # 添加匹配分数列
            age_filtered_df['匹配分数'] = age_filtered_df['标签'].apply(calculate_match_score)
            
            # 筛选有匹配标签的动作
            filtered_df = age_filtered_df[age_filtered_df['匹配分数'] > 0]
            
            if filtered_df.empty:
                messagebox.showwarning("警告", f"在动作库中未找到适合{age}岁且匹配所选标签的训练动作")
                return
            
            # 按匹配分数降序排序
            filtered_df = filtered_df.sort_values(by='匹配分数', ascending=False)
            
            # 将动作转换为记录列表
            available_actions = filtered_df.to_dict('records')
            
            # 每节课3个动作，共12节课，需要36个动作
            required_actions = 12 * 3
            
            # 如果动作不足以组成12节课（每节3个训练项目），则允许重复
            if len(available_actions) < required_actions:
                # 直接创建12节课，每节课3个不重复的动作
                grouped_actions = []
                
                for lesson in range(12):
                    # 为当前课随机选择3个不重复的动作
                    current_lesson_actions = []
                    # 复制可用动作列表，避免修改原始列表
                    lesson_action_pool = available_actions.copy()
                    
                    # 如果可用动作少于3个，重新补充动作池
                    if len(lesson_action_pool) < 3:
                        lesson_action_pool = available_actions.copy()
                    
                    # 选择3个不重复的动作
                    for _ in range(3):
                        # 随机选择一个动作
                        action_index = random.randrange(len(lesson_action_pool))
                        current_lesson_actions.append(lesson_action_pool.pop(action_index))
                        
                        # 如果用完了所有动作，重新补充动作池
                        if not lesson_action_pool and _ < 2:  # 还需要选择更多动作
                            lesson_action_pool = [a for a in available_actions if a not in current_lesson_actions]
                    
                    grouped_actions.append(current_lesson_actions)
            else:
                # 有足够的不重复动作，随机打乱后每3个分组
                random.shuffle(available_actions)
                grouped_actions = []
                
                # 先确保使用不同的动作
                actions_copy = available_actions.copy()
                for lesson in range(12):
                    lesson_actions = []
                    for _ in range(3):
                        action = actions_copy.pop(0)
                        lesson_actions.append(action)
                    grouped_actions.append(lesson_actions)
            
            # 创建输出目录
            try:
                os.makedirs(output_dir, exist_ok=True)
            except PermissionError:
                messagebox.showerror("错误", f"没有权限创建输出目录: {output_dir}\n请尝试使用管理员权限运行或选择其他目录。")
                return
            
            # 生成Word文档
            doc = Document()
            
            # 设置默认字体为微软雅黑
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            
            # 通过XML直接设置文档默认字体
            try:
                default_fonts = doc._part.document.body.xpath('//w:rFonts')
                for font in default_fonts:
                    font.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
            except:
                pass
                
            # 设置文档级别默认字体
            try:
                doc._element.get_or_add_docDefaults().get_or_add_rPrDefault().get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
            except:
                pass
            
            # 设置所有标题字体为微软雅黑
            self.set_doc_fonts(doc)
            
            # 清除所有段落布局
            self.clear_paragraph_formatting(doc)
            
            # 设置文档标题
            title = doc.add_heading(f"{name} 的{len(grouped_actions)}节课感统训练方案", level=0)
            # 设置标题字体
            for run in title.runs:
                self.set_run_font(run)
            
            # 添加基本信息 - 姓名年龄和生成日期
            info_para = doc.add_paragraph()
            name_run = info_para.add_run(f"姓名: {name}    ")
            name_run.bold = True
            self.set_run_font(name_run)
                
            age_run = info_para.add_run(f"年龄: {age} 岁    ")
            age_run.bold = True
            self.set_run_font(age_run)
                
            date_run = info_para.add_run(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d')}")
            date_run.bold = True
            self.set_run_font(date_run)
            
            # 单独一段显示训练标签
            tag_para = doc.add_paragraph()
            tag_run = tag_para.add_run(f"训练标签: {', '.join(selected_tags)}")
            tag_run.bold = True
            self.set_run_font(tag_run)
            
            # 为了空间，添加一个空行
            doc.add_paragraph()
            
            # 逐课添加训练内容
            for i, lesson_actions in enumerate(grouped_actions, 1):
                # 添加课程标题 - 使用蓝色文字
                lesson_title = doc.add_heading(f"第{i}节课:", level=2)
                for run in lesson_title.runs:
                    run.font.color.rgb = docx.shared.RGBColor(65, 105, 225)  # 皇家蓝
                    self.set_run_font(run)
                
                # 为每节课创建一个段落，包含所有动作
                for j, action in enumerate(lesson_actions, 1):
                    # 动作名称和描述在同一段落中
                    # 如果是第一个动作，创建新段落；如果不是，在前一个动作后继续添加内容
                    if j == 1:
                        action_para = doc.add_paragraph()
                    
                    # 动作名称
                    if j > 1:
                        # 如果不是第一个动作，先添加换行符分隔上一个动作
                        action_para.add_run("\n").bold = False
                        self.set_run_font(action_para.runs[-1])
                    
                    action_name_run = action_para.add_run(f"动作名称 {j}: ")
                    action_name_run.bold = True
                    self.set_run_font(action_name_run)
                    
                    action_name_content = action_para.add_run(f"{action['动作名称']}")
                    self.set_run_font(action_name_content)
                    
                    # 换行
                    newline_run = action_para.add_run("\n")
                    self.set_run_font(newline_run)
                    
                    # 训练目标（使用动作标签，去掉#，用、分隔）
                    action_target_run = action_para.add_run(f"训练目标: ")
                    action_target_run.bold = True
                    self.set_run_font(action_target_run)
                    
                    # 处理标签：将#分隔改为、分隔
                    if isinstance(action['标签'], str):
                        tags = action['标签'].replace('#', '、')
                        # 如果第一个字符是、，则去掉
                        if tags.startswith('、'):
                            tags = tags[1:]
                    else:
                        tags = ""
                    
                    action_target_content = action_para.add_run(f"{tags}")
                    self.set_run_font(action_target_content)
                    
                    # 换行
                    newline_run = action_para.add_run("\n")
                    self.set_run_font(newline_run)
                    
                    # 动作描述（不使用斜体）
                    action_desc_run = action_para.add_run(f"动作描述: ")
                    self.set_run_font(action_desc_run)
                    
                    action_desc_content = action_para.add_run(f"{action['动作描述']}")
                    self.set_run_font(action_desc_content)
                
                # 在每节课之间添加一个空行，除非是最后一节课
                if i < len(grouped_actions):
                    doc.add_paragraph()
            
            # 保存文档
            output_filename = os.path.join(
                output_dir,
                f"{name}_训练方案_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            
            try:
                doc.save(output_filename)
                messagebox.showinfo("成功", f"训练方案已生成: {output_filename}")
                # 尝试打开文件所在的文件夹
                try:
                    os.startfile(os.path.dirname(output_filename))
                except:
                    pass  # 如果无法打开文件夹，则忽略错误
            except PermissionError:
                messagebox.showerror("错误", f"无法保存文件: {output_filename}\n请检查文件是否被其他程序占用或没有写入权限。")
            
        except FileNotFoundError as e:
            messagebox.showerror("错误", f"找不到文件: {str(e)}")
        except PermissionError:
            messagebox.showerror("错误", "没有足够的权限读取或写入文件")
        except Exception as e:
            messagebox.showerror("错误", f"生成训练方案失败: {str(e)}")
            raise e  # 重新抛出异常，以便调用者可以捕获具体错误
    
    def clear_paragraph_formatting(self, doc):
        """清除所有段落的布局设置"""
        try:
            # 尝试通过直接操作OOXML实现
            # 找到所有段落
            paragraphs = doc._element.xpath('//w:p')
            
            for paragraph in paragraphs:
                # 查找段落属性节点
                pPr = paragraph.xpath('./w:pPr')
                if pPr:
                    # 删除段落属性中与布局相关的节点
                    for pPr_elem in pPr:
                        # 删除对齐方式设置
                        jc_elements = pPr_elem.xpath('./w:jc')
                        for jc in jc_elements:
                            pPr_elem.remove(jc)
                        
                        # 删除缩进设置
                        ind_elements = pPr_elem.xpath('./w:ind')
                        for ind in ind_elements:
                            pPr_elem.remove(ind)
                        
                        # 删除段落间距设置
                        spacing_elements = pPr_elem.xpath('./w:spacing')
                        for spacing in spacing_elements:
                            pPr_elem.remove(spacing)
                        
                        # 删除分页设置
                        pageBreak_elements = pPr_elem.xpath('./w:pageBreakBefore')
                        for pageBreak in pageBreak_elements:
                            pPr_elem.remove(pageBreak)
                        
                        # 删除制表位设置
                        tabs_elements = pPr_elem.xpath('./w:tabs')
                        for tabs in tabs_elements:
                            pPr_elem.remove(tabs)
                        
                        # 删除边框设置
                        border_elements = pPr_elem.xpath('./w:pBdr')
                        for border in border_elements:
                            pPr_elem.remove(border)
                        
                        # 删除其他可能的段落格式设置
                        # 段落底纹设置
                        shd_elements = pPr_elem.xpath('./w:shd')
                        for shd in shd_elements:
                            pPr_elem.remove(shd)
                            
                        # 段落调整类型
                        adjustRightInd_elements = pPr_elem.xpath('./w:adjustRightInd')
                        for adjustRightInd in adjustRightInd_elements:
                            pPr_elem.remove(adjustRightInd)
                            
                        # 自动断行设置
                        wordWrap_elements = pPr_elem.xpath('./w:wordWrap')
                        for wordWrap in wordWrap_elements:
                            pPr_elem.remove(wordWrap)
            
            # 修改python-docx对象中的段落属性
            for paragraph in doc.paragraphs:
                # 重置段落格式
                paragraph.paragraph_format.alignment = None
                paragraph.paragraph_format.left_indent = None
                paragraph.paragraph_format.right_indent = None
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing = None
                paragraph.paragraph_format.space_before = None
                paragraph.paragraph_format.space_after = None
                paragraph.paragraph_format.keep_together = False
                paragraph.paragraph_format.keep_with_next = False
                paragraph.paragraph_format.page_break_before = False
                paragraph.paragraph_format.widow_control = False
        except Exception as e:
            # 如果出现错误，忽略但不中断流程
            pass
    
    def set_doc_fonts(self, doc):
        """设置文档中所有字体为微软雅黑"""
        # 设置标题字体为微软雅黑
        for i in range(0, 10):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = '微软雅黑'
                # 确保所有字体设置都使用微软雅黑
                if hasattr(heading_style.font, 'ascii_font'):
                    heading_style.font.ascii_font.name = '微软雅黑'
                if hasattr(heading_style.font, 'eastasia_font'):
                    heading_style.font.eastasia_font.name = '微软雅黑'
                if hasattr(heading_style.font, 'cs_font'):
                    heading_style.font.cs_font.name = '微软雅黑'
            except:
                pass
        
        # 确保Normal样式的所有字体变体都使用微软雅黑
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        if hasattr(style.font, 'ascii_font'):
            style.font.ascii_font.name = '微软雅黑'
        if hasattr(style.font, 'eastasia_font'):
            style.font.eastasia_font.name = '微软雅黑'
        if hasattr(style.font, 'cs_font'):
            style.font.cs_font.name = '微软雅黑'
    
    def set_run_font(self, run):
        """设置运行块的所有字体变体为微软雅黑"""
        run.font.name = '微软雅黑'
        # 设置更多字体属性来确保中文正确显示
        if hasattr(run.font, 'ascii_font'):
            run.font.ascii_font.name = '微软雅黑'
        if hasattr(run.font, 'eastasia_font'):
            run.font.eastasia_font.name = '微软雅黑'
        if hasattr(run.font, 'cs_font'):
            run.font.cs_font.name = '微软雅黑'
        # 添加XML字体设置尝试
        try:
            run._element.get_or_add_rPr().get_or_add_rFonts().set('eastAsia', '微软雅黑')
        except:
            pass 