import os
import pandas as pd
import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import docx
import logging
from flask import current_app
import io

# 设置日志记录
logger = logging.getLogger(__name__)

class WebTrainingPlanGenerator:
    """网页版训练方案生成器，基于原有的训练方案生成器功能修改"""
    
    def __init__(self):
        """初始化训练方案生成器"""
        pass
    
    def get_age_range(self, age):
        """根据年龄获取对应的年龄范围分类
        
        Args:
            age: 学生年龄
            
        Returns:
            str: 年龄范围分类，如 '4-6', '7-9', '10-12'
        """
        try:
            age = int(age)
            if 4 <= age <= 6:
                return '4-6'
            elif 7 <= age <= 9:
                return '7-9'
            elif 10 <= age <= 12:
                return '10-12'
            else:
                return None
        except ValueError:
            return None
            
    def generate_training_plan(self, name, age, selected_tags, action_db_path):
        """生成训练方案，并返回Word文档的二进制流
        
        Args:
            name: 学生姓名
            age: 学生年龄
            selected_tags: 选中的训练标签列表
            action_db_path: 动作库Excel文件路径
            
        Returns:
            bytes: Word文档的二进制数据
        """
        try:
            # 确定年龄范围
            age_range = self.get_age_range(age)
            if not age_range:
                logger.warning(f"无法确定年龄范围，学生{name}年龄为{age}，不在4-12岁之间")
                return None, "无法确定年龄范围，请确保年龄在4-12岁之间"
            
            # 读取动作库
            logger.info(f"正在读取动作库: {action_db_path}")
            df = pd.read_excel(action_db_path)
            
            # 过滤掉动作名称为空的行
            df = df.dropna(subset=['动作名称']).reset_index(drop=True)
            
            # 筛选年龄范围符合的动作
            age_filtered_df = df[df['年龄范围'] == age_range]
            
            if age_filtered_df.empty:
                logger.warning(f"在动作库中未找到适合{age}岁的训练动作")
                return None, f"在动作库中未找到适合{age}岁的训练动作"
            
            # 计算标签匹配分数，优先选择标签匹配度高的动作
            def calculate_match_score(tag_str):
                if not isinstance(tag_str, str):
                    return 0
                
                action_tags = [tag.strip() for tag in tag_str.split('#') if tag.strip()]
                if not action_tags:
                    return 0
                
                # 计算用户选择标签与动作标签的交集
                matching_tags = set(action_tags).intersection(set(selected_tags))
                
                # 计算匹配分数
                # 1. 完全匹配情况：用户选择标签与动作标签完全相同
                if set(action_tags) == set(selected_tags):
                    return 100
                
                # 2. 用户只选了一个标签，动作也只有这一个标签
                if len(selected_tags) == 1 and len(action_tags) == 1 and action_tags[0] in selected_tags:
                    return 90
                
                # 3. 匹配标签占动作标签的比例
                action_match_ratio = len(matching_tags) / len(action_tags)
                
                # 4. 匹配标签占用户选择标签的比例
                user_match_ratio = len(matching_tags) / len(selected_tags)
                
                # 综合评分 (加权平均)
                return (action_match_ratio * 30) + (user_match_ratio * 30) + (len(matching_tags) * 10)
            
            # 添加匹配分数列
            age_filtered_df['匹配分数'] = age_filtered_df['标签'].apply(calculate_match_score)
            
            # 筛选有匹配标签的动作
            filtered_df = age_filtered_df[age_filtered_df['匹配分数'] > 0]
            
            if filtered_df.empty:
                logger.warning(f"在动作库中未找到适合{age}岁且匹配所选标签的训练动作")
                return None, f"在动作库中未找到适合{age}岁且匹配所选标签的训练动作"
            
            # 按匹配分数降序排序
            filtered_df = filtered_df.sort_values(by='匹配分数', ascending=False)
            
            # 将动作转换为记录列表
            available_actions = filtered_df.to_dict('records')
            
            # 每节课3个动作，共12节课，需要36个动作
            required_actions = 12 * 3
            
            # 如果动作不足以组成12节课（每节3个训练项目），则允许重复
            if len(available_actions) < required_actions:
                # 直接创建12节课，每节课3个不重复的动作
                grouped_actions = []
                
                for lesson in range(12):
                    # 为当前课随机选择3个不重复的动作
                    current_lesson_actions = []
                    # 复制可用动作列表，避免修改原始列表
                    lesson_action_pool = available_actions.copy()
                    
                    # 如果可用动作少于3个，重新补充动作池
                    if len(lesson_action_pool) < 3:
                        lesson_action_pool = available_actions.copy()
                    
                    # 选择3个不重复的动作
                    for _ in range(3):
                        # 随机选择一个动作
                        action_index = random.randrange(len(lesson_action_pool))
                        current_lesson_actions.append(lesson_action_pool.pop(action_index))
                        
                        # 如果用完了所有动作，重新补充动作池
                        if not lesson_action_pool and _ < 2:  # 还需要选择更多动作
                            lesson_action_pool = [a for a in available_actions if a not in current_lesson_actions]
                    
                    grouped_actions.append(current_lesson_actions)
            else:
                # 有足够的不重复动作，随机打乱后每3个分组
                random.shuffle(available_actions)
                grouped_actions = []
                
                # 先确保使用不同的动作
                actions_copy = available_actions.copy()
                for lesson in range(12):
                    lesson_actions = []
                    for _ in range(3):
                        action = actions_copy.pop(0)
                        lesson_actions.append(action)
                    grouped_actions.append(lesson_actions)
            
            # 生成Word文档
            doc = Document()
            
            # 设置默认字体为微软雅黑
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            
            # 通过XML直接设置文档默认字体
            try:
                default_fonts = doc._part.document.body.xpath('//w:rFonts')
                for font in default_fonts:
                    font.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
            except:
                pass
                
            # 设置文档级别默认字体
            try:
                doc._element.get_or_add_docDefaults().get_or_add_rPrDefault().get_or_add_rPr().get_or_add_rFonts().set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
            except:
                pass
            
            # 设置所有标题字体为微软雅黑
            self.set_doc_fonts(doc)
            
            # 清除所有段落布局
            self.clear_paragraph_formatting(doc)
            
            # 设置文档标题
            title = doc.add_heading(f"{name} 的{len(grouped_actions)}节课感统训练方案", level=0)
            # 设置标题字体
            for run in title.runs:
                self.set_run_font(run)
            
            # 添加基本信息 - 姓名年龄和生成日期
            info_para = doc.add_paragraph()
            name_run = info_para.add_run(f"姓名: {name}    ")
            name_run.bold = True
            self.set_run_font(name_run)
                
            age_run = info_para.add_run(f"年龄: {age} 岁    ")
            age_run.bold = True
            self.set_run_font(age_run)
                
            date_run = info_para.add_run(f"生成日期: {datetime.datetime.now().strftime('%Y-%m-%d')}")
            date_run.bold = True
            self.set_run_font(date_run)
            
            # 添加所选标签信息
            tag_para = doc.add_paragraph()
            tag_title_run = tag_para.add_run("训练标签: ")
            tag_title_run.bold = True
            self.set_run_font(tag_title_run)
            
            tag_content_run = tag_para.add_run(", ".join(selected_tags))
            self.set_run_font(tag_content_run)
            
            doc.add_paragraph()  # 空行
            
            # 添加所有课程内容
            for lesson_index, lesson_actions in enumerate(grouped_actions, 1):
                # 添加课次标题
                lesson_header = doc.add_heading(f"第{lesson_index}课", level=1)
                for run in lesson_header.runs:
                    self.set_run_font(run)
                    run.bold = True
                
                # 为每节课创建所有动作
                for action_index, action in enumerate(lesson_actions, 1):
                    self.add_action_item(doc, action_index, action)
                
                # 在每个课次后添加空行（最后一个课次除外）
                if lesson_index < len(grouped_actions):
                    doc.add_paragraph()
            
            # 保存文档到内存中
            document_stream = io.BytesIO()
            doc.save(document_stream)
            document_stream.seek(0)
            
            logger.info(f"成功为{name}生成训练方案")
            return document_stream.getvalue(), None
            
        except Exception as e:
            logger.exception(f"生成训练方案失败: {str(e)}")
            return None, f"生成训练方案失败: {str(e)}"
    
    def clear_paragraph_formatting(self, doc):
        """清除文档中段落的格式设置"""
        try:
            # 尝试通过直接操作OOXML实现
            # 找到所有段落
            paragraphs = doc._element.xpath('//w:p')
            
            for paragraph in paragraphs:
                # 查找段落属性节点
                pPr = paragraph.xpath('./w:pPr')
                if pPr:
                    # 删除段落属性中与布局相关的节点
                    for pPr_elem in pPr:
                        # 删除对齐方式设置
                        jc_elements = pPr_elem.xpath('./w:jc')
                        for jc in jc_elements:
                            pPr_elem.remove(jc)
                        
                        # 删除缩进设置
                        ind_elements = pPr_elem.xpath('./w:ind')
                        for ind in ind_elements:
                            pPr_elem.remove(ind)
                        
                        # 删除段落间距设置
                        spacing_elements = pPr_elem.xpath('./w:spacing')
                        for spacing in spacing_elements:
                            pPr_elem.remove(spacing)
                        
                        # 删除分页设置
                        pageBreak_elements = pPr_elem.xpath('./w:pageBreakBefore')
                        for pageBreak in pageBreak_elements:
                            pPr_elem.remove(pageBreak)
                        
                        # 删除制表位设置
                        tabs_elements = pPr_elem.xpath('./w:tabs')
                        for tabs in tabs_elements:
                            pPr_elem.remove(tabs)
                        
                        # 删除边框设置
                        border_elements = pPr_elem.xpath('./w:pBdr')
                        for border in border_elements:
                            pPr_elem.remove(border)
                        
                        # 删除其他可能的段落格式设置
                        # 段落底纹设置
                        shd_elements = pPr_elem.xpath('./w:shd')
                        for shd in shd_elements:
                            pPr_elem.remove(shd)
                            
                        # 段落调整类型
                        adjustRightInd_elements = pPr_elem.xpath('./w:adjustRightInd')
                        for adjustRightInd in adjustRightInd_elements:
                            pPr_elem.remove(adjustRightInd)
                            
                        # 自动断行设置
                        wordWrap_elements = pPr_elem.xpath('./w:wordWrap')
                        for wordWrap in wordWrap_elements:
                            pPr_elem.remove(wordWrap)
            
            # 修改python-docx对象中的段落属性
            for paragraph in doc.paragraphs:
                # 重置段落格式
                paragraph.paragraph_format.alignment = None
                paragraph.paragraph_format.left_indent = None
                paragraph.paragraph_format.right_indent = None
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing = None
                paragraph.paragraph_format.space_before = None
                paragraph.paragraph_format.space_after = None
                paragraph.paragraph_format.keep_together = False
                paragraph.paragraph_format.keep_with_next = False
                paragraph.paragraph_format.page_break_before = False
                paragraph.paragraph_format.widow_control = False
        except Exception as e:
            # 如果出现错误，使用备用方法
            try:
                # 获取并修改第一个段落的样式
                first_style = doc.paragraphs[0].style
                
                # 清除段落格式
                for paragraph in doc.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.style = first_style
                    
                    # 清除段落中的格式设置
                    for run in paragraph.runs:
                        # 保留粗体和字体设置
                        is_bold = run.bold
                        font_name = run.font.name
                        
                        # 重置运行的格式
                        run.font.reset()
                        
                        # 恢复所需的格式
                        run.bold = is_bold
                        if font_name:
                            run.font.name = font_name
                            
                # 尝试清除所有段落的缩进和间距
                for paragraph in doc.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = 0
                    paragraph_format.right_indent = 0
                    paragraph_format.first_line_indent = 0
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph_format.space_before = 0
                    paragraph_format.space_after = 0
            except:
                pass
    
    def set_doc_fonts(self, doc):
        """设置文档中所有字体为微软雅黑"""
        # 设置标题字体为微软雅黑
        for i in range(0, 10):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = '微软雅黑'
                # 确保所有字体设置都使用微软雅黑
                if hasattr(heading_style.font, 'ascii_font'):
                    heading_style.font.ascii_font.name = '微软雅黑'
                if hasattr(heading_style.font, 'eastasia_font'):
                    heading_style.font.eastasia_font.name = '微软雅黑'
                if hasattr(heading_style.font, 'cs_font'):
                    heading_style.font.cs_font.name = '微软雅黑'
            except:
                pass
        
        # 确保Normal样式的所有字体变体都使用微软雅黑
        try:
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'
            if hasattr(style.font, 'ascii_font'):
                style.font.ascii_font.name = '微软雅黑'
            if hasattr(style.font, 'eastasia_font'):
                style.font.eastasia_font.name = '微软雅黑'
            if hasattr(style.font, 'cs_font'):
                style.font.cs_font.name = '微软雅黑'
        except:
            pass
        
        # 设置所有段落的字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                self.set_run_font(run)
    
    def set_run_font(self, run):
        """设置运行块的所有字体变体为微软雅黑"""
        run.font.name = '微软雅黑'
        # 设置更多字体属性来确保中文正确显示
        if hasattr(run.font, 'ascii_font'):
            run.font.ascii_font.name = '微软雅黑'
        if hasattr(run.font, 'eastasia_font'):
            run.font.eastasia_font.name = '微软雅黑'
        if hasattr(run.font, 'cs_font'):
            run.font.cs_font.name = '微软雅黑'
        # 添加XML字体设置尝试
        try:
            run._element.get_or_add_rPr().get_or_add_rFonts().set('eastAsia', '微软雅黑')
        except:
            pass
    
    def add_lesson_header(self, doc, lesson_index):
        """添加课次标题
        
        Args:
            doc: Word文档对象
            lesson_index: 课次索引(从1开始)
        """
        lesson_header = doc.add_heading(f"第{lesson_index}课", level=1)
        for run in lesson_header.runs:
            self.set_run_font(run)
            run.bold = True
    
    def add_action_item(self, doc, action_index, action):
        """添加动作项
        
        Args:
            doc: Word文档对象
            action_index: 动作在当前课次中的索引(从1开始)
            action: 动作信息字典
        """
        # 添加动作名称和其他信息在同一段落
        action_para = doc.add_paragraph()
        action_name_run = action_para.add_run(f"动作名称 {action_index}: ")
        action_name_run.bold = True
        self.set_run_font(action_name_run)
        
        action_name_content = action_para.add_run(f"{action.get('动作名称', '未命名动作')}")
        self.set_run_font(action_name_content)
        
        # 换行
        newline_run = action_para.add_run("\n")
        self.set_run_font(newline_run)
        
        # 训练目标（使用动作标签，去掉#，用、分隔）
        action_target_run = action_para.add_run(f"训练目标: ")
        action_target_run.bold = True
        self.set_run_font(action_target_run)
        
        # 处理标签：将#分隔改为、分隔
        if isinstance(action.get('标签'), str):
            tags = action.get('标签').replace('#', '、')
            # 如果第一个字符是、，则去掉
            if tags.startswith('、'):
                tags = tags[1:]
        else:
            tags = ""
        
        action_target_content = action_para.add_run(f"{tags}")
        self.set_run_font(action_target_content)
        
        # 换行
        newline_run = action_para.add_run("\n")
        self.set_run_font(newline_run)
        
        # 动作描述
        if action.get('动作描述'):
            action_desc_run = action_para.add_run(f"动作描述: ")
            action_desc_run.bold = True
            self.set_run_font(action_desc_run)
            
            action_desc_content = action_para.add_run(f"{action.get('动作描述', '')}")
            self.set_run_font(action_desc_content)
        
        # 动作材料（添加到同一个段落）
        if action.get('材料'):
            # 换行
            newline_run = action_para.add_run("\n")
            self.set_run_font(newline_run)
            
            material_title = action_para.add_run("所需材料: ")
            material_title.bold = True
            self.set_run_font(material_title)
            
            material_content = action_para.add_run(f"{action.get('材料', '')}")
            self.set_run_font(material_content)
        
        # 训练提示（添加到同一个段落）
        if action.get('训练提示'):
            # 换行
            newline_run = action_para.add_run("\n")
            self.set_run_font(newline_run)
            
            tips_title = action_para.add_run("训练提示: ")
            tips_title.bold = True
            self.set_run_font(tips_title)
            
            tips_content = action_para.add_run(f"{action.get('训练提示', '')}")
            self.set_run_font(tips_content) 